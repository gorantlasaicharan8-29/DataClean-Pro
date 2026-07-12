"""Report generation – PDF (FPDF2), DOCX (python-docx), HTML (Jinja2)."""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any

import pandas as pd
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jinja2 import Template

from app.services.statistics_service import get_full_statistics
from app.services.insights_service import generate_insights


# ─────────────────────────────────────────────────────────────────────────────
# PDF
# ─────────────────────────────────────────────────────────────────────────────

class _ReportPDF(FPDF):
    """Custom FPDF subclass with header / footer branding."""

    def header(self) -> None:
        self.set_font("Helvetica", "B", 10)
        self.set_text_color(100, 100, 100)
        self.cell(0, 8, "DataClean Pro – Data Analysis Report", align="L", new_x="LMARGIN", new_y="NEXT")
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(4)

    def footer(self) -> None:
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")


def generate_pdf_report(
    df: pd.DataFrame,
    original_df: pd.DataFrame,
    insights: dict[str, Any] | None = None,
    stats: dict[str, Any] | None = None,
    charts_data: list[Any] | None = None,
    filename: str = "report",
) -> bytes:
    if stats is None:
        stats = get_full_statistics(df)
    if insights is None:
        insights = generate_insights(df)

    pdf = _ReportPDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)

    # ── Title page ────────────────────────────────────────────────────
    pdf.add_page()
    pdf.ln(40)
    pdf.set_font("Helvetica", "B", 28)
    pdf.set_text_color(41, 98, 255)
    pdf.cell(0, 15, "DataClean Pro", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 16)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 10, "Data Analysis Report", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(10)
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, f"File: {filename}", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 8, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 8, f"Original: {len(original_df)} rows  |  Cleaned: {len(df)} rows", align="C", new_x="LMARGIN", new_y="NEXT")

    # ── Dataset summary ───────────────────────────────────────────────
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(41, 98, 255)
    pdf.cell(0, 10, "1. Dataset Summary", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(0, 0, 0)

    summary_items = [
        ("Rows (original)", str(len(original_df))),
        ("Rows (cleaned)", str(len(df))),
        ("Columns", str(len(df.columns))),
        ("Missing values (cleaned)", str(int(df.isna().sum().sum()))),
        ("Duplicates (cleaned)", str(int(df.duplicated().sum()))),
    ]
    for label, val in summary_items:
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(60, 7, label + ":", new_x="RIGHT")
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 7, val, new_x="LMARGIN", new_y="NEXT")

    # ── Cleaning comparison ───────────────────────────────────────────
    pdf.ln(6)
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(41, 98, 255)
    pdf.cell(0, 10, "2. Cleaning Summary", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(0, 0, 0)

    orig_miss = int(original_df.isna().sum().sum())
    clean_miss = int(df.isna().sum().sum())
    orig_dups = int(original_df.duplicated().sum())
    clean_dups = int(df.duplicated().sum())

    compare = [
        ("Metric", "Before", "After"),
        ("Rows", str(len(original_df)), str(len(df))),
        ("Columns", str(len(original_df.columns)), str(len(df.columns))),
        ("Missing values", str(orig_miss), str(clean_miss)),
        ("Duplicates", str(orig_dups), str(clean_dups)),
    ]
    col_w = 60
    for row_idx, row in enumerate(compare):
        if row_idx == 0:
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_fill_color(41, 98, 255)
            pdf.set_text_color(255, 255, 255)
        else:
            pdf.set_font("Helvetica", "", 10)
            pdf.set_fill_color(245, 245, 245) if row_idx % 2 == 0 else pdf.set_fill_color(255, 255, 255)
            pdf.set_text_color(0, 0, 0)
        for cell in row:
            pdf.cell(col_w, 7, cell, border=1, fill=True)
        pdf.ln()

    pdf.set_text_color(0, 0, 0)

    # ── Statistics ────────────────────────────────────────────────────
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(41, 98, 255)
    pdf.cell(0, 10, "3. Statistical Analysis", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_text_color(0, 0, 0)

    for col_stat in stats.get("columns", [])[:15]:
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(0, 7, f"Column: {col_stat['name']} ({col_stat['dtype']})", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 9)
        detail_keys = ["mean", "median", "std", "min", "max", "missing", "unique"]
        parts = []
        for k in detail_keys:
            v = col_stat.get(k)
            if v is not None:
                parts.append(f"{k}: {v}")
        if parts:
            pdf.cell(0, 6, "  " + " | ".join(parts), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

    # ── Insights ──────────────────────────────────────────────────────
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(41, 98, 255)
    pdf.cell(0, 10, "4. AI Insights", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 10)

    for item in insights.get("insights", []):
        icon = item.get("icon", "")
        msg = item.get("message", "")
        # FPDF can struggle with emoji; fall back gracefully
        try:
            pdf.multi_cell(0, 6, f"{icon}  {msg}", new_x="LMARGIN", new_y="NEXT")
        except Exception:
            pdf.multi_cell(0, 6, f"- {msg}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)

    pdf.ln(4)
    pdf.set_font("Helvetica", "I", 10)
    summary_text = insights.get("summary", "")
    pdf.multi_cell(0, 6, summary_text)

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# DOCX
# ─────────────────────────────────────────────────────────────────────────────

def generate_docx_report(
    df: pd.DataFrame,
    original_df: pd.DataFrame,
    insights: dict[str, Any] | None = None,
    stats: dict[str, Any] | None = None,
    filename: str = "report",
) -> bytes:
    if stats is None:
        stats = get_full_statistics(df)
    if insights is None:
        insights = generate_insights(df)

    doc = Document()

    # Title
    title = doc.add_heading("DataClean Pro – Data Analysis Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"File: {filename}\n"
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        f"Original: {len(original_df)} rows  |  Cleaned: {len(df)} rows"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Dataset summary
    doc.add_heading("1. Dataset Summary", level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = "Light Shading Accent 1"
    data = [
        ("Rows (original)", str(len(original_df))),
        ("Rows (cleaned)", str(len(df))),
        ("Columns", str(len(df.columns))),
        ("Missing values", str(int(df.isna().sum().sum()))),
        ("Duplicates", str(int(df.duplicated().sum()))),
    ]
    for i, (label, val) in enumerate(data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = val

    # Cleaning comparison
    doc.add_heading("2. Cleaning Summary", level=1)
    table = doc.add_table(rows=5, cols=3)
    table.style = "Light Shading Accent 1"
    headers = ["Metric", "Before", "After"]
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    compare_data = [
        ("Rows", str(len(original_df)), str(len(df))),
        ("Columns", str(len(original_df.columns)), str(len(df.columns))),
        ("Missing", str(int(original_df.isna().sum().sum())), str(int(df.isna().sum().sum()))),
        ("Duplicates", str(int(original_df.duplicated().sum())), str(int(df.duplicated().sum()))),
    ]
    for i, (m, b, a) in enumerate(compare_data, start=1):
        table.rows[i].cells[0].text = m
        table.rows[i].cells[1].text = b
        table.rows[i].cells[2].text = a

    # Statistics
    doc.add_heading("3. Statistical Analysis", level=1)
    for col_stat in stats.get("columns", [])[:15]:
        doc.add_heading(f"{col_stat['name']} ({col_stat['dtype']})", level=3)
        detail_keys = ["mean", "median", "std", "min", "max", "missing", "unique"]
        parts = [f"{k}: {col_stat[k]}" for k in detail_keys if col_stat.get(k) is not None]
        if parts:
            doc.add_paragraph(" | ".join(parts))

    # Insights
    doc.add_heading("4. AI Insights", level=1)
    for item in insights.get("insights", []):
        doc.add_paragraph(f"{item.get('icon', '-')}  {item.get('message', '')}", style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(insights.get("summary", ""))
    run.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# HTML
# ─────────────────────────────────────────────────────────────────────────────

_HTML_TEMPLATE = Template("""\
<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>DataClean Pro Report – {{ filename }}</title>
<style>
  :root{--primary:#2962ff;--bg:#f9fafb;--card:#fff;--text:#1e293b;--muted:#64748b}
  *{box-sizing:border-box;margin:0;padding:0}
  body{font-family:'Segoe UI',system-ui,sans-serif;background:var(--bg);color:var(--text);line-height:1.6;padding:2rem}
  .container{max-width:900px;margin:auto}
  h1{color:var(--primary);margin-bottom:.5rem}
  h2{color:var(--primary);margin:2rem 0 1rem;border-bottom:2px solid var(--primary);padding-bottom:.3rem}
  h3{margin:1rem 0 .5rem}
  .meta{color:var(--muted);margin-bottom:2rem}
  table{width:100%;border-collapse:collapse;margin:1rem 0}
  th,td{padding:.5rem .75rem;border:1px solid #e2e8f0;text-align:left}
  th{background:var(--primary);color:#fff}
  tr:nth-child(even){background:#f1f5f9}
  .insight{padding:.5rem 0;border-bottom:1px solid #e2e8f0}
  .summary{background:#eef2ff;padding:1rem;border-radius:.5rem;margin-top:1rem;font-style:italic}
  .badge{display:inline-block;padding:2px 8px;border-radius:4px;font-size:.75rem;font-weight:600;margin-left:.5rem}
  .badge.info{background:#dbeafe;color:#1e40af}
  .badge.warning{background:#fef3c7;color:#92400e}
  .badge.critical{background:#fee2e2;color:#991b1b}
</style>
</head>
<body>
<div class="container">
  <h1>DataClean Pro</h1>
  <p class="meta">
    File: {{ filename }} &nbsp;|&nbsp;
    Generated: {{ generated }} &nbsp;|&nbsp;
    Original: {{ orig_rows }} rows &nbsp;|&nbsp;
    Cleaned: {{ clean_rows }} rows
  </p>

  <h2>1. Dataset Summary</h2>
  <table>
    <tr><th>Metric</th><th>Value</th></tr>
    <tr><td>Rows (original)</td><td>{{ orig_rows }}</td></tr>
    <tr><td>Rows (cleaned)</td><td>{{ clean_rows }}</td></tr>
    <tr><td>Columns</td><td>{{ columns }}</td></tr>
    <tr><td>Missing values</td><td>{{ missing }}</td></tr>
    <tr><td>Duplicates</td><td>{{ duplicates }}</td></tr>
  </table>

  <h2>2. Cleaning Summary</h2>
  <table>
    <tr><th>Metric</th><th>Before</th><th>After</th></tr>
    <tr><td>Rows</td><td>{{ orig_rows }}</td><td>{{ clean_rows }}</td></tr>
    <tr><td>Columns</td><td>{{ orig_cols }}</td><td>{{ columns }}</td></tr>
    <tr><td>Missing</td><td>{{ orig_missing }}</td><td>{{ missing }}</td></tr>
    <tr><td>Duplicates</td><td>{{ orig_dups }}</td><td>{{ duplicates }}</td></tr>
  </table>

  <h2>3. Statistical Analysis</h2>
  {% for col in col_stats %}
  <h3>{{ col.name }} <small>({{ col.dtype }})</small></h3>
  <table>
    <tr>
    {% for key in stat_keys %}
      {% if col[key] is not none %}
      <th>{{ key }}</th>
      {% endif %}
    {% endfor %}
    </tr>
    <tr>
    {% for key in stat_keys %}
      {% if col[key] is not none %}
      <td>{{ col[key] }}</td>
      {% endif %}
    {% endfor %}
    </tr>
  </table>
  {% endfor %}

  <h2>4. AI Insights</h2>
  {% for item in insights %}
  <div class="insight">
    {{ item.icon }} {{ item.message }}
    <span class="badge {{ item.severity }}">{{ item.severity }}</span>
  </div>
  {% endfor %}

  <div class="summary">{{ summary }}</div>
</div>
</body>
</html>
""")


def generate_html_report(
    df: pd.DataFrame,
    original_df: pd.DataFrame,
    insights: dict[str, Any] | None = None,
    stats: dict[str, Any] | None = None,
    filename: str = "report",
) -> str:
    if stats is None:
        stats = get_full_statistics(df)
    if insights is None:
        insights = generate_insights(df)

    return _HTML_TEMPLATE.render(
        filename=filename,
        generated=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        orig_rows=len(original_df),
        clean_rows=len(df),
        columns=len(df.columns),
        orig_cols=len(original_df.columns),
        missing=int(df.isna().sum().sum()),
        orig_missing=int(original_df.isna().sum().sum()),
        duplicates=int(df.duplicated().sum()),
        orig_dups=int(original_df.duplicated().sum()),
        col_stats=stats.get("columns", [])[:15],
        stat_keys=["mean", "median", "std", "min", "max", "missing", "unique"],
        insights=insights.get("insights", []),
        summary=insights.get("summary", ""),
    )
